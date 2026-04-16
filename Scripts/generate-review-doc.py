#!/usr/bin/env python3
"""Generate a .docx compliance review document for manager/legal review."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re, sys

ROOT = Path(__file__).parent.parent

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def build_doc():
    doc = Document()

    # -- Adjust default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ======================================================================
    # TITLE PAGE
    # ======================================================================
    doc.add_paragraph()  # spacer
    title = doc.add_heading("WorkLogger", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Data Protection & Compliance Review\nfor Legal / Compliance Assessment")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("April 2026\nPrepared for PwC Compliance & Legal Team")
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc,
        "This document provides a comprehensive overview of the WorkLogger application, "
        "its data processing activities, privacy controls, GDPR compliance measures, "
        "and risk assessment — intended for review by PwC's Legal and Compliance teams.",
        size=11)

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS (manual)
    # ======================================================================
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Application Overview",
        "3. Data Processing Activities",
        "4. Data Categories & Collection",
        "5. Privacy Controls & Settings",
        "6. GDPR Compliance Measures",
        "7. Data Subject Rights (Art. 15–22)",
        "8. Security Measures (Art. 32)",
        "9. Risk Assessment & Mitigations",
        "10. Test Coverage & Quality Assurance",
        "11. Deployment & Distribution",
        "12. Recommendations for Legal/Compliance",
        "Appendix A — Record of Processing Activities (ROPA)",
        "Appendix B — Data Protection Impact Assessment (DPIA)",
    ]
    for item in toc_items:
        add_para(doc, item, size=11)

    doc.add_page_break()

    # ======================================================================
    # 1. EXECUTIVE SUMMARY
    # ======================================================================
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "WorkLogger is a lightweight macOS menu bar application developed for personal work activity tracking. "
        "It logs application usage, window titles, Safari tab changes, idle periods, and manual task entries "
        "as local JSONL files. A Python-based report generator produces weekly Excel timesheets enriched with "
        "git commits from configured repositories.")
    add_para(doc,
        "Key compliance characteristics:")
    bullets = [
        "All data is processed and stored exclusively on the user's local Mac — no network transmission, no cloud storage, no third-party services",
        "GDPR compliance implemented: first-launch consent dialog, auto-purge (90-day default), domain-only URL logging, file permissions (chmod 600/700), data export and erasure from menu bar",
        "Privacy by design and default (Art. 25): most restrictive settings enabled out of the box",
        "Full GDPR documentation: Privacy Notice (Art. 13), ROPA (Art. 30), DPIA (Art. 35)",
        "Comprehensive automated test coverage (113 tests) gating all builds — no untested code ships",
        "Open-source codebase available for audit at https://github.com/Ralo93/tracker",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    # ======================================================================
    # 2. APPLICATION OVERVIEW
    # ======================================================================
    add_heading(doc, "2. Application Overview", level=1)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Application name", "WorkLogger"],
            ["Platform", "macOS (13+ Ventura or later)"],
            ["Language", "Swift (app) + Python (report generator)"],
            ["Distribution", "Built from source via Swift Package Manager"],
            ["Code signing", "Ad-hoc (self-signed) — no Apple Developer ID"],
            ["Network access", "None — zero network calls"],
            ["Source code", "https://github.com/Ralo93/tracker"],
            ["License", "Internal / to be determined"],
        ])

    add_para(doc, "")
    add_heading(doc, "How it works", level=2)
    add_para(doc,
        "WorkLogger runs as a menu bar icon (\"WL\") and monitors application switches, VS Code project changes, "
        "Safari tab changes, and idle/lock/sleep events. Each event is appended to a daily JSONL log file. "
        "Users can add manual entries via a global hotkey (Cmd+Shift+L). "
        "A weekly Excel report can be generated from the menu bar or command line.")

    # ======================================================================
    # 3. DATA PROCESSING ACTIVITIES
    # ======================================================================
    add_heading(doc, "3. Data Processing Activities", level=1)
    add_table(doc,
        ["Activity", "Lawful Basis", "Justification"],
        [
            ["Track active applications and window titles", "Legitimate interest (Art. 6(1)(f))", "User's own interest in accurate personal time tracking"],
            ["Log Safari tab activity", "Legitimate interest + Consent", "Optional feature, configurable via toggles; explicit consent on first launch"],
            ["Detect idle periods, screen lock, sleep/wake", "Legitimate interest (Art. 6(1)(f))", "Accurately delimit work sessions"],
            ["Store manual task entries", "Consent (Art. 6(1)(a))", "User voluntarily creates entries"],
            ["Read git commit history", "Legitimate interest (Art. 6(1)(f))", "Enrich weekly time report with work output"],
            ["Generate weekly Excel report", "Legitimate interest (Art. 6(1)(f))", "Purpose of the tool — produce a time tracking report"],
        ])

    add_para(doc, "")
    add_para(doc,
        "Note on consent in employment context: If an employer or manager requires team members to use WorkLogger, "
        "consent may not be freely given (GDPR Recital 43). In such cases, legitimate interest of the employee "
        "(accurate self-reporting) is the primary lawful basis. A Legitimate Interest Assessment (LIA) should be "
        "documented by the organization.",
        italic=True, size=10)

    # ======================================================================
    # 4. DATA CATEGORIES
    # ======================================================================
    add_heading(doc, "4. Data Categories & Collection", level=1)
    add_heading(doc, "What IS collected", level=2)
    add_table(doc,
        ["Data Type", "Source", "Purpose"],
        [
            ["Active application name", "macOS workspace notifications", "Track which apps are used and when"],
            ["Window titles", "CGWindowListCopyWindowInfo", "Identify what the user is working on"],
            ["VS Code project name", "Window title parsing", "Track project-level time in VS Code"],
            ["Safari tab names", "AppleScript automation", "Track web-based work activities"],
            ["Safari URLs", "AppleScript automation", "Identify visited sites (domain only by default)"],
            ["Idle start/end", "CGEventSource idle time", "Detect breaks and inactive periods"],
            ["Screen lock/unlock", "Distributed notifications", "Detect break periods"],
            ["Sleep/wake", "Workspace notifications", "Detect system sleep periods"],
            ["Manual entries", "User input", "User-created task descriptions with time and duration"],
            ["Git commits", "Local git repos (report only)", "Enrich report with commit details"],
        ])

    add_para(doc, "")
    add_heading(doc, "What is NOT collected", level=2)
    not_collected = [
        "Screen content, screenshots, or pixel data",
        "Keystrokes or keyboard input (beyond idle detection)",
        "Clipboard contents",
        "Network traffic or browsing history beyond the active Safari tab",
        "Audio, video, or microphone data",
        "Data from other users on shared Macs",
    ]
    for item in not_collected:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    # ======================================================================
    # 5. PRIVACY CONTROLS
    # ======================================================================
    add_heading(doc, "5. Privacy Controls & Settings", level=1)
    add_para(doc,
        "All privacy settings are accessible via Preferences → General tab and are persisted in the config file. "
        "Each setting has a tooltip explaining its exact effect.")
    add_table(doc,
        ["Setting", "Default", "Effect"],
        [
            ["safariTrackingEnabled", "true", "Set to false to completely disable Safari monitoring"],
            ["logSafariURLs", "true", "Set to false to log tab names but not URLs"],
            ["logSafariDomainOnly", "true", "Only store domain (e.g. github.com, not the full path)"],
            ["retentionDays", "90", "Days to keep log files before automatic deletion"],
            ["skipApps", "(list)", "App names excluded from report descriptions"],
            ["skipSafariExact", "(list)", "Safari tab titles excluded by exact match"],
            ["skipSafariContains", "(list)", "Safari tab titles excluded by substring match"],
            ["showSafariTimeInReport", "false", "Show accumulated time per Safari tab in the report"],
        ])

    add_para(doc, "")
    add_heading(doc, "URL Sanitization", level=2)
    add_para(doc,
        "By default (logSafariDomainOnly: true), URLs are reduced to the domain only:")
    add_para(doc,
        "https://portal.client.com/project/12345/docs?token=abc → portal.client.com",
        italic=True, size=10)
    add_para(doc,
        "OAuth tokens, authorization codes, session IDs, and similar sensitive URL parameters are never persisted. "
        "Query strings and fragments are always stripped, even when domain-only mode is disabled.")

    # ======================================================================
    # 6. GDPR COMPLIANCE MEASURES
    # ======================================================================
    add_heading(doc, "6. GDPR Compliance Measures", level=1)

    measures = [
        ("First-launch consent (Art. 7)", "On first launch, a consent dialog lists exactly what is tracked and requires explicit \"I Agree\" before any logging begins. Declining exits the app immediately."),
        ("Privacy notice (Art. 13/14)", "PRIVACY.md provides a full Art. 13 privacy notice covering lawful basis, data categories, rights, retention, security measures, and controller template."),
        ("Auto-purge / retention (Art. 5(1)(e))", "Logs older than the retention period (default 90 days) are automatically deleted on every app launch."),
        ("Data minimization (Art. 5(1)(c))", "Domain-only URL logging, 200-character string truncation, configurable feature toggles to disable unnecessary collection."),
        ("Privacy by design (Art. 25)", "Most restrictive settings enabled by default. No data leaves the device. All file permissions are owner-only."),
        ("Documentation (Art. 30, 35)", "Record of Processing Activities (ROPA.md) and Data Protection Impact Assessment (DPIA.md) provided."),
        ("No third-party data sharing", "Zero network calls. No analytics, telemetry, cloud storage, or third-party services."),
    ]
    for title, desc in measures:
        add_heading(doc, title, level=2)
        add_para(doc, desc)

    # ======================================================================
    # 7. DATA SUBJECT RIGHTS
    # ======================================================================
    add_heading(doc, "7. Data Subject Rights (Art. 15–22)", level=1)
    add_table(doc,
        ["Right", "Article", "Implementation"],
        [
            ["Right of access", "Art. 15", "\"Export All My Data…\" menu item creates a zip archive of all JSONL logs and config"],
            ["Right to rectification", "Art. 16", "Retroactive Quick Log tab allows adding corrected entries; JSONL files editable in any text editor"],
            ["Right to erasure", "Art. 17", "\"Delete All My Data…\" menu item with double confirmation; auto-purge of old logs"],
            ["Right to restriction", "Art. 18", "Safari toggle, skipApps config, quit app to stop processing"],
            ["Right to portability", "Art. 20", "JSONL is open format; zip export includes all data"],
            ["Right to object", "Art. 21", "Quit app, decline consent, or delete all data"],
            ["Withdraw consent", "Art. 7(3)", "Delete config to reset consent; quit app at any time"],
        ])

    # ======================================================================
    # 8. SECURITY MEASURES
    # ======================================================================
    add_heading(doc, "8. Security Measures (Art. 32)", level=1)
    add_table(doc,
        ["Measure", "Implementation"],
        [
            ["File permissions", "Log files: 0600 (owner read/write only), log directory: 0700, config: 0600"],
            ["Encryption at rest", "Relies on macOS FileVault (full-disk encryption, standard on managed corporate Macs)"],
            ["Data minimization", "URLs logged as domain-only by default; all strings truncated to 200 characters"],
            ["No network transmission", "Zero network calls — no analytics, telemetry, or cloud sync"],
            ["Automatic purging", "Logs older than retention period deleted on every launch"],
            ["Access control", "Single-user tool; no multi-user access, no shared accounts"],
        ])

    # ======================================================================
    # 9. RISK ASSESSMENT
    # ======================================================================
    add_heading(doc, "9. Risk Assessment & Mitigations", level=1)
    add_para(doc,
        "The following table summarizes all identified data protection and security concerns, their risk levels, "
        "and current mitigation status.")
    add_table(doc,
        ["#", "Concern", "Risk", "Status"],
        [
            ["1", "Data at rest — no encryption", "Medium", "Mitigated (chmod 600, FileVault documented)"],
            ["2", "Sensitive data in URLs", "High", "Resolved (domain-only default, query stripping, toggles)"],
            ["3", "No data retention / purge", "Medium", "Resolved (90-day auto-purge on launch)"],
            ["4", "No access controls", "Low-Med", "Mitigated (chmod 600/700)"],
            ["5", "Safari tracking — privacy", "Medium", "Resolved (toggles, domain-only default)"],
            ["6", "Window title leakage", "Medium", "Mitigated (documented, consent, skipApps)"],
            ["7", "No user consent mechanism", "High", "Resolved (first-launch dialog, menu item)"],
            ["8", "Ad-hoc code signing", "Medium", "Documented"],
            ["9", "LaunchAgent auto-start", "Low", "Documented"],
            ["10", "Screen Recording & Accessibility", "Medium", "Documented"],
            ["11", "Git commit data in reports", "Low", "Documented"],
            ["12", "No audit trail", "Low", "Documented"],
            ["13", "Data portability & deletion", "Low", "Resolved (export zip, delete all, auto-purge)"],
            ["14", "GDPR rights implementation", "Medium", "Resolved (Art. 15–22 all implemented)"],
            ["15", "GDPR documentation", "Medium", "Resolved (PRIVACY.md, ROPA.md, DPIA.md)"],
            ["16", "Automated test coverage", "Medium", "Resolved (113 tests, build-gated)"],
        ])

    add_para(doc, "")
    add_heading(doc, "Residual risks", level=2)
    residual = [
        "Data at rest not encrypted at application level — mitigated by FileVault recommendation",
        "Window titles may contain confidential document names — mitigated by consent and skipApps filter",
        "Ad-hoc code signing may be blocked by corporate MDM — requires IT approval or Developer ID ($99/year)",
        "No application-level audit trail — acceptable for single-user personal tool",
    ]
    for item in residual:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    # ======================================================================
    # 10. TEST COVERAGE
    # ======================================================================
    add_heading(doc, "10. Test Coverage & Quality Assurance", level=1)
    add_para(doc,
        "All compliance-critical controls are verified by automated tests that must pass before any build can be produced.")
    add_table(doc,
        ["Suite", "Tests", "Coverage area"],
        [
            ["Swift — URL sanitization", "12+", "Domain-only extraction, query stripping, token removal, IP/unicode/encoded URLs"],
            ["Swift — File permissions", "3", "chmod 600 on log files, retroactive entries, persistence across writes"],
            ["Swift — Auto-purge", "7", "Retention periods, boundary, non-JSONL preservation, edge cases"],
            ["Swift — Config & consent", "11", "Privacy field encode/decode, consent mechanism, Safari toggles, defaults"],
            ["Swift — Data minimization", "8+", "String truncation (200 chars), privacy-by-default verification"],
            ["Python — Report pipeline", "46", "Smart aggregation, Teams filtering, description builder, manual entries"],
        ])
    add_para(doc, "")
    add_para(doc,
        "Build gate: make app → make build → make test. Both Swift and Python suites must pass before a binary is produced. "
        "Total: 113 tests across 21 suites.",
        bold=True)

    # ======================================================================
    # 11. DEPLOYMENT
    # ======================================================================
    add_heading(doc, "11. Deployment & Distribution", level=1)
    add_para(doc,
        "WorkLogger is distributed as source code via GitHub. Users build the .app bundle locally using "
        "make app, which requires only the Xcode Command Line Tools (free) and Python 3.10+.")
    add_para(doc, "")
    add_heading(doc, "Prerequisites for end users", level=2)
    prereqs = [
        "macOS 13 (Ventura) or later",
        "Xcode Command Line Tools (xcode-select --install)",
        "Python 3.10+ with openpyxl (pip3 install openpyxl)",
        "Grant Accessibility, Screen Recording, and Safari Automation permissions on first launch",
    ]
    for item in prereqs:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    add_para(doc, "")
    add_heading(doc, "macOS permissions explained", level=2)
    add_table(doc,
        ["Permission", "Reason", "Scope"],
        [
            ["Accessibility", "Read window titles via CGWindowList; global Cmd+Shift+L hotkey", "Only reads window names, not content"],
            ["Screen Recording", "Required by macOS for CGWindowListCopyWindowInfo", "Only reads window titles, not screen pixels"],
            ["Automation (Safari)", "Read the active tab name and URL", "Only reads current tab, not browsing history"],
        ])

    # ======================================================================
    # 12. RECOMMENDATIONS
    # ======================================================================
    add_heading(doc, "12. Recommendations for Legal/Compliance", level=1)

    recs = [
        ("For immediate deployment (personal use)", [
            "The application can be used immediately for personal time tracking",
            "All high-risk concerns have been resolved (consent, URL sanitization, data retention)",
            "Users should ensure FileVault is enabled on their Mac",
            "The source code is available for security audit at https://github.com/Ralo93/tracker",
        ]),
        ("For team-wide deployment", [
            "Complete the controller details in PRIVACY.md and ROPA.md templates",
            "Conduct a Legitimate Interest Assessment (LIA) if usage is mandatory",
            "Coordinate with IT to whitelist the app in MDM profiles (Accessibility, Screen Recording)",
            "Consider obtaining an Apple Developer ID ($99/year) for proper code signing",
            "Brief team members on what data is collected and their rights",
        ]),
        ("For formal compliance sign-off", [
            "Review the attached DPIA (Appendix B) and complete the assessor fields",
            "Consult the DPO on the ROPA entry (Appendix A)",
            "Determine whether Works Council (Betriebsrat) consultation is required under BetrVG §87",
            "Verify alignment with PwC's internal data classification and handling policies",
        ]),
    ]
    for title, items in recs:
        add_heading(doc, title, level=2)
        for item in items:
            p = doc.add_paragraph(item, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_page_break()

    # ======================================================================
    # APPENDIX A — ROPA
    # ======================================================================
    add_heading(doc, "Appendix A — Record of Processing Activities (ROPA)", level=1)
    add_para(doc, "GDPR Article 30 — Record of Processing Activities", italic=True)
    add_para(doc, "")

    ropa_text = (ROOT / "ROPA.md").read_text()
    for line in ropa_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("---"):
            continue
        if line.startswith("# "):
            continue  # skip main title, we have our own heading
        if line.startswith("## "):
            add_heading(doc, line.lstrip("# "), level=2)
        elif line.startswith("*"):
            add_para(doc, line.strip("*").strip(), italic=True, size=10)
        elif line.startswith("| ") and "---" not in line:
            # Skip markdown tables — they're better recreated above
            add_para(doc, line.replace("|", "").strip(), size=10)
        else:
            add_para(doc, line, size=11)

    doc.add_page_break()

    # ======================================================================
    # APPENDIX B — DPIA
    # ======================================================================
    add_heading(doc, "Appendix B — Data Protection Impact Assessment (DPIA)", level=1)
    add_para(doc, "GDPR Article 35 — Data Protection Impact Assessment", italic=True)
    add_para(doc, "")

    dpia_text = (ROOT / "DPIA.md").read_text()
    for line in dpia_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("---"):
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            add_heading(doc, line.lstrip("# "), level=2)
        elif line.startswith("### "):
            add_heading(doc, line.lstrip("# "), level=3)
        elif line.startswith("*"):
            add_para(doc, line.strip("*").strip(), italic=True, size=10)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)
        elif line.startswith("| ") and "---" not in line:
            add_para(doc, line.replace("|", "  ").strip(), size=10)
        else:
            add_para(doc, line, size=11)

    # ======================================================================
    # SAVE
    # ======================================================================
    out = ROOT / "WorkLogger_Compliance_Review.docx"
    doc.save(str(out))
    print(f"✓ Saved to {out}")
    return out

if __name__ == "__main__":
    build_doc()
